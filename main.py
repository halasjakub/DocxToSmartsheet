import os
import json
import time
import pyautogui
import logging
import smartsheet
import config
import smartsheet_data
import tkinter
from tkinter import Label, Button
from tkinter import messagebox, filedialog
from docx import Document
from collections import defaultdict
from datetime import datetime

# Constants
VENDOR_DATA_FILE = "vendor_data.json"
LOG_FILE = "rwsheet.log"

# Initialize Smartsheet client
smartsheet_client = smartsheet.Smartsheet(config.API_KEY)
logging.basicConfig(filename=LOG_FILE, level=logging.INFO)

# Main GUI window
gui_main_window = tkinter.Tk()
gui_main_window.title("DocxToSmartsheet")
gui_main_window.geometry(f"{pyautogui.size()[0]}x{pyautogui.size()[1]}")


def get_d_code_en_num_from_docx(file_path):
    """Extract the D code and Enhancement number from a DOCX file."""
    document = Document(file_path)

    keywords = ["Client UAT for Task ", "Enhancement Number "]

    d_code = None
    e_number = "-"

    for section in document.sections:
        header = section.header

        # Check if the header has at least two paragraphs
        if len(header.paragraphs) >= 2:
            second_paragraph = header.paragraphs[1].text

            if keywords[0].lower() in second_paragraph.lower():
                start_index = second_paragraph.lower().find(keywords[0]) + len(keywords[0])
                d_code = second_paragraph[start_index:].strip()

        # Check if the header has at least three paragraphs
        if len(header.paragraphs) >= 3:
            third_paragraph = header.paragraphs[2].text

            if third_paragraph != "":
                if keywords[1].lower() in third_paragraph.lower():
                    start_index = third_paragraph.lower().find(keywords[1]) + len(keywords[1])
                    e_number = third_paragraph[start_index:].strip()

    return d_code, e_number


def get_count_import_docx_data_to_smartsheet(file_path):
    """Process DOCX data and update Smartsheet."""
    processing_time_start = time.time()

    d_code, e_number = get_d_code_en_num_from_docx(file_path)
    document = Document(file_path)

    date_format = "%d-%b-%Y"  # dd-mmm-yyyy
    columns = [0, 1, 4, 8, 10]
    table = document.tables[0]

    issue_type_counts = defaultdict(int)
    issue_delay_counts = 0
    checked_rows = 0

    for row in table.rows[2:]:
        row_data = [
            row.cells[col_idx].text.strip() for col_idx in columns if col_idx < len(row.cells)
        ]

        if len(set(row_data)) == 1:
            continue
        else:
            uat_start_date = datetime.strptime(row.cells[1].text.strip(), date_format)
            uat_end_date = datetime.strptime(row.cells[8].text.strip(), date_format)
            uat_time = (uat_end_date - uat_start_date).days
            issue_type = row.cells[10].text.strip()

            if d_code and issue_type:
                issue_type_counts[d_code, issue_type] += 1

            if uat_time > 2:
                print(row.cells[0].text.strip(), f" Yes: {uat_time}")
                issue_delay_counts += 1
            else:
                print(row.cells[0].text.strip(), f" No: {uat_time}")

        checked_rows += 1

    print(f"No. system issues resolved in > 2 business days: {issue_delay_counts}")

    issues_delay_number_column_id = searching_column_id(
        vendor_url, smartsheet_data.column_name[5]
    )

    row_id = searching_row_id(
        vendor_url, smartsheet_data.column_name[0], d_code, smartsheet_data.column_name[3], e_number
    )

    for smartsheet_column_name in smartsheet_data.column_name[7:]:
        column_id = searching_column_id(vendor_url, smartsheet_column_name)
        update_single_cell_in_smartsheet(vendor_url, row_id, column_id, 0)

    print("\n D-code | Enhancement number | Issues number | Issue type")

    for (d_code, issue_type), issues_number in issue_type_counts.items():
        print(f"{d_code} | {e_number} | {issues_number} | {issue_type}")
        column_id = searching_column_id(vendor_url, issue_type)
        update_single_cell_in_smartsheet(vendor_url, row_id, column_id, issues_number)

    update_single_cell_in_smartsheet(
        vendor_url, row_id, issues_delay_number_column_id, issue_delay_counts
    )

    processing_time_end = time.time()
    processing_time = round(processing_time_end - processing_time_start)

    tkinter.messagebox.showinfo(
        "Process completed",
        f"{checked_rows} updated records in {processing_time}s.\n"
        f"No. system issues resolved in > 2 business days: {issue_delay_counts}"
    )


def searching_row_id(sheet_id, d_code_column_name, d_code, en_num_column_name, e_number):
    """Search for a row ID based on d_code and e_number in specified columns."""
    
    
    def check_values_in_row(row):
        """Check if the row contains the correct d_code and e_number."""
        found_d_code = False
        found_en_num = False

        for cell in row.cells:
            if cell.column_id == column_names.get(d_code_column_name):
                found_d_code = cell.value
            if cell.column_id == column_names.get(en_num_column_name):
                found_en_num = cell.value

        return found_d_code == d_code and found_en_num == e_number

    try:
        sheet = smartsheet_client.Sheets.get_sheet(sheet_id)
        if hasattr(sheet, 'columns'):
            column_names = {col.title: col.id for col in sheet.columns}
            row_id = None

            for row in sheet.rows:
                if check_values_in_row(row):
                    row_id = row.id
                    return row_id

            return row_id

        else:
            tkinter.messagebox.showinfo(
                "Searching row ID",
                f"Invalid or not found sheet: {sheet}"
            )
            return None
    
    except smartsheet.exceptions.ApiError as e:
        tkinter.messagebox.showinfo(
                "Searching row ID",
                f"Error during API call: {e}"
            )
        return None    
    

def searching_column_id(sheet_id, new_value):
    """Search for the column ID by matching column title with new_value."""
    try:
        sheet = smartsheet_client.Sheets.get_sheet(sheet_id)
        if hasattr(sheet, 'columns'):
            column_id = None

            for column in sheet.columns:
                if column.title == new_value:
                    column_id = column.id
                    return column_id
        else:
            tkinter.messagebox.showinfo(
                "Searching column ID",
                f"Invalid or not found sheet: {sheet}"
            )
            return None
    
    except smartsheet.exceptions.ApiError as e:
        tkinter.messagebox.showinfo(
                "Searching column ID",
                f"Error during API call: {e}"
            )
        return None


def update_single_cell_in_smartsheet(sheet_id, row_id, column_id, new_value):
    """Update a single cell in a row on a Smartsheet."""
    try:
        sheet = smartsheet_client.Sheets.get_sheet(sheet_id)

        updated_cell = smartsheet.models.Cell()
        updated_cell.column_id = column_id
        updated_cell.value = new_value
        updated_cell.strict = False

        updated_row = smartsheet.models.Row()
        updated_row.id = row_id
        updated_row.cells = [updated_cell]
    
        if isinstance(sheet, smartsheet.models.Sheet):
            response = smartsheet_client.Sheets.update_rows(sheet.id, [updated_row])
            if response.message == 'SUCCESS':
                print("The data has been changed")
            else:
                print(f"Error: {response.message}")
        else:
            tkinter.messagebox.showinfo(
                "Update single cell in Smartsheet",
                f"Invalid or not found sheet: {sheet}"
            )
            return None
    
    except smartsheet.exceptions.ApiError as e:
        tkinter.messagebox.showinfo(
                "Update single cell in Smartsheet",
                f"Error during API call: {e}"
            )
        return None


def read_vendor_data():
    """Open and read the vendor data file."""
    with open(VENDOR_DATA_FILE, 'r', encoding='utf-8') as openfile:
        return json.load(openfile)


def refresh_label(vendor_name, vendor_url, word_file):
    """Refresh the label displaying vendor name, vendor URL, and word file."""
    label_vendor = Label(gui_main_window, text=f"{vendor_name}")
    label_vendor.place(x=530, y=45)
    label_file = Label(gui_main_window, text=f"{word_file}")
    label_file.place(x=170, y=135)
    gui_main_window.update()


def ask_about_word():
    """Prompt the user to select a Word file."""
    global vendor_name, vendor_url, word_file
    word_file = filedialog.askopenfilename()
    
    # Check if the file is a DOCX file and if it exists
    if not word_file.lower().endswith('.docx') or not os.path.isfile(word_file):
        tkinter.messagebox.showinfo(
            "Open",
            f"Invalid file format or file not found: {word_file}"
        )
        return None

    if word_file and os.path.exists(word_file):
        ask_about_file = tkinter.messagebox.askquestion("Open", f"Is {word_file} a selected file?")

        if ask_about_file == 'yes':
            # Only refresh the label once vendor and file are both selected
            if vendor_name != "" and vendor_url != "":
                refresh_label(vendor_name, vendor_url, word_file)
            else:
                refresh_label("", "", word_file)  # Only file selected, show file info


def set_vendor_url(vendor_name_local, vendor_url_local, word_file):
    """Set the vendor URL and refresh the label if both vendor and file are selected."""
    global vendor_name, vendor_url
    vendor_name = vendor_name_local  # Now assign values to them
    vendor_url = vendor_url_local

    if word_file != "":  # If both vendor and file are selected
        refresh_label(vendor_name, vendor_url, word_file)
    else:  # Only update vendor info if file hasn't been selected yet
        refresh_label(vendor_name, vendor_url, "")


def connection_tester():
    """Test the Smartsheet connection."""
    try:
        account_info = smartsheet_client.Users.get_current_user()
        tkinter.messagebox.showinfo(
            "Connection test",
            f"Connection test completed, connected by: {account_info.email}"
        )
    except smartsheet.exceptions.ApiError as e:
        tkinter.messagebox.showinfo(
            "Connection test",
            f"Connection test completed, error: {e}"
        )


def starter():
    """Start the process by checking the vendor URL and Word file."""
    global vendor_url, word_file, vendor_name

    if vendor_url != "" and word_file != "":
        ask_about_start = tkinter.messagebox.askquestion(
            "Start",
            f"Do you want to start the process for: \nVendor: {vendor_name}\n"
            f"File path: {word_file}\n?"
        )
        if ask_about_start == 'yes':
            print(f"Vendor name: {vendor_name}")
            print(f"Vendor URL: {vendor_url}")
            print(f"File path: {word_file}")
            d_code, e_number = get_d_code_en_num_from_docx(word_file)
            print(f"d_code: {d_code}\ne_number: {e_number}")
            get_count_import_docx_data_to_smartsheet(word_file)


# Read vendor data from JSON file
vendor_data = read_vendor_data()

# Initialize vendor variables
word_file = ""  # Initialize an empty string for the word file path
vendor_name = ""  # Initialize an empty string for the vendor name
vendor_url = ""  # Initialize an empty string for the vendor URL

# Title labels
title_label_vendor = Label(gui_main_window, text="1. Select Vendor", font=(20))
title_label_vendor.place(x=10, y=10)

# Create vendor buttons dynamically
for index, (name, url) in enumerate(vendor_data.items()):
    Button(
        gui_main_window,
        text=name,
        height=1,
        width=12,
        command=lambda vendor_name=name, vendor_url=url: set_vendor_url(vendor_name, vendor_url, word_file)
    ).place(x=40 + (index * 120), y=40)

# File upload section
title_label_file = Label(gui_main_window, text="2. Upload file", font=(20))
title_label_file.place(x=10, y=100)

file_button = Button(
    gui_main_window,
    text="Upload",
    height=1,
    width=12,
    command=ask_about_word
)
file_button.place(x=40, y=130)

# Start process section
title_label_start = Label(gui_main_window, text="3. Start process", font=(20))
title_label_start.place(x=10, y=190)

start_process = Button(
    gui_main_window,
    text="Start",
    height=1,
    width=12,
    command=starter
)
start_process.place(x=40, y=220)

# Connection test section
connection_test = Button(
    gui_main_window,
    text="Connection test",
    height=1,
    width=24,
    command=connection_tester
)
connection_test.place(x=40, y=280)

# Start the main loop of the Tkinter application
gui_main_window.mainloop()
